"""Generate the TMS Training Management Policy & SOP as a Word (.docx) document.

Run:  python docs/build_policy_sop.py
Output: docs/TMS_Policy_and_SOP.docx

Content is grounded in the live TMS logic (roles, annual mandate, data-integrity
guards, the training cycle, verification, effectiveness, and the calculation
basis — including the conducted-only Summary rule and the exit policy). Keep this
script as the source of truth; regenerate the .docx after policy changes.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x1A, 0x7A, 0x3C)   # green
ACCENT = RGBColor(0x2E, 0x6D, 0xA4)  # blue
GREY = RGBColor(0x6E, 0x6E, 0x73)

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'TMS_Policy_and_SOP.docx')

doc = Document()

# ── Base styles ───────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz, col in [('Heading 1', 15, BRAND), ('Heading 2', 12.5, ACCENT), ('Heading 3', 11, RGBColor(0x1D,0x1D,0x1F))]:
    st = doc.styles[lvl]
    st.font.name = 'Calibri'
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True


def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def table(headers, rows, widths=None, header_bg='1A7A3C'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)
        _set_cell_bg(hdr[i], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def bullets(items, style='List Bullet'):
    for it in items:
        doc.add_paragraph(it, style=style)


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style='List Number')


def footer_pagenum():
    sec = doc.sections[0]
    f = sec.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = f.add_run('TMS Policy & SOP · Balrampur Chini Mills Ltd · Page ')
    run.font.size = Pt(8); run.font.color.rgb = GREY
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    f._p.append(fld)


# ── TITLE PAGE ────────────────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TRAINING MANAGEMENT SYSTEM')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = BRAND
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Policy & Standard Operating Procedure')
r.font.size = Pt(16); r.font.color.rgb = ACCENT
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Learning & Development')
r.font.size = Pt(12); r.font.color.rgb = GREY
for _ in range(8):
    doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line, bold in [('Balrampur Chini Mills Ltd.', True),
                   ('Document No.: TMS-POL-SOP-001   |   Version: 1.0', False),
                   ('Classification: Internal   |   Owner: Corporate L&D', False),
                   ('Effective Date: 07 June 2026', False)]:
    rr = meta.add_run(line + '\n'); rr.bold = bold; rr.font.size = Pt(11 if bold else 9.5)
    if not bold:
        rr.font.color.rgb = GREY
doc.add_page_break()

# ── 0. DOCUMENT CONTROL ───────────────────────────────────────────────────────
doc.add_heading('Document Control', level=1)
table(['Field', 'Detail'],
      [['Document Title', 'TMS Training Management Policy & Standard Operating Procedure'],
       ['Document No.', 'TMS-POL-SOP-001'],
       ['Version', '1.0'],
       ['Effective Date', '07 June 2026'],
       ['Owner', 'Corporate Learning & Development'],
       ['Approved By', 'Head — Human Resources'],
       ['Review Cycle', 'Annual, or on material change to the training process'],
       ['Applies To', 'All manufacturing units / plants and Corporate L&D']],
      widths=[1.8, 4.6])
doc.add_heading('Revision History', level=2)
table(['Version', 'Date', 'Author', 'Summary of Change'],
      [['1.0', '07 Jun 2026', 'Corporate L&D', 'Initial issue — policy + end-to-end SOP for the Training Management System.']],
      widths=[0.9, 1.1, 1.6, 2.8])
doc.add_page_break()

# ── 1. PURPOSE & SCOPE ────────────────────────────────────────────────────────
doc.add_heading('1. Purpose & Scope', level=1)
doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph(
    'This document defines the policy and the standard operating procedure (SOP) for managing '
    'employee training through the Training Management System (TMS). It establishes how training '
    'need is identified, planned, conducted, recorded, verified, measured for effectiveness, and '
    'reported — so that every plant can demonstrate, on demand and with traceable evidence, that '
    'the mandated training was delivered.')
doc.add_heading('1.2 Scope', level=2)
doc.add_paragraph(
    'The SOP covers the complete training cycle for all employees (Blue Collared and White '
    'Collared) across every plant, plus centrally-hosted corporate programmes. It governs data '
    'entry, the QR-based attendance and feedback capture, the monthly summary and compliance '
    'dashboard, and the export of statutory/management records.')
doc.add_heading('1.3 Objective', level=2)
bullets([
    'Meet the annual man-hour training mandate for every employee.',
    'Maintain a single, trustworthy source of truth for training data across all plants.',
    'Be audit-ready at all times — every record traceable to the individual, date, and approver.',
    'Measure not just attendance but the effectiveness of training.'])
doc.add_page_break()

# ── 2. DEFINITIONS ────────────────────────────────────────────────────────────
doc.add_heading('2. Definitions & Abbreviations', level=1)
table(['Term', 'Meaning'],
      [['TNI', 'Training Need Identification — the nominated demand: one row per employee per programme.'],
       ['Programme Master', 'The canonical, approved list of programme names per plant. Acts as the gatekeeper for all programme selection.'],
       ['Calendar (2C-plan)', 'Planned sessions, each with an auto-generated Session Code.'],
       ['2C / Programme Details', 'Conducted-session actuals (dates, faculty, hours, cost, feedback).'],
       ['2A / Training Records', 'Individual attendance — one row per person per session.'],
       ['Session Code', 'Unique session identifier: UNIT/TYPE/NNN/FY/Bnn (e.g. BCM/TEC/001/26-27/B01).'],
       ['BC / WC', 'Blue Collared / White Collared employee.'],
       ['Coverage %', 'Of the employees nominated for a programme in TNI, the proportion who were trained.'],
       ['Compliance %', 'Man-hours delivered as a proportion of the annual man-hour mandate.'],
       ['SPOC', 'Single Point of Contact — the plant-level training administrator.'],
       ['Central L&D', 'Corporate Learning & Development — cross-plant oversight and verification.'],
       ['QR Attendance/Feedback', 'Phone-scannable QR that captures attendance or feedback with no app/login.']],
      widths=[1.6, 4.8])
doc.add_page_break()

# ── 3. GOVERNANCE & ROLES ─────────────────────────────────────────────────────
doc.add_heading('3. Governance & Roles', level=1)
doc.add_paragraph('The TMS enforces three roles. Access is least-privilege and role-gated.')
table(['Role', 'Who', 'Responsibilities', 'System Rights'],
      [['SPOC', 'Plant training administrator',
        'Maintain employee master; upload TNI; plan calendar; generate QR; record 2C/2A; review monthly summary.',
        'Read/write scoped to own plant only.'],
       ['Central L&D', 'Corporate L&D team',
        'Cross-plant oversight; verify conducted sessions; plan corporate programmes; monitor compliance.',
        'Read across all plants; verify/approve; manage central calendar. 2FA mandatory.'],
       ['Admin', 'System administrator',
        'User & access management; impersonate a plant for support; configuration; backups.',
        'Full access; can act on behalf of any plant. 2FA mandatory.']],
      widths=[0.9, 1.5, 2.6, 1.6])
doc.add_heading('3.1 RACI — Training Cycle', level=2)
table(['Activity', 'SPOC', 'Central L&D', 'Admin'],
      [['Employee master upkeep', 'R/A', 'I', 'C'],
       ['TNI upload & confirm', 'R/A', 'I', 'C'],
       ['Calendar planning', 'R/A', 'C', 'I'],
       ['Conduct & QR attendance', 'R/A', 'I', 'I'],
       ['2C / 2A recording', 'R/A', 'C', 'I'],
       ['Verification of sessions', 'C', 'R/A', 'I'],
       ['Effectiveness review', 'R', 'A', 'I'],
       ['Compliance reporting', 'R', 'A', 'I'],
       ['Configuration & access', 'I', 'C', 'R/A']],
      widths=[2.6, 1.2, 1.4, 1.2])
doc.add_paragraph('R = Responsible · A = Accountable · C = Consulted · I = Informed', style='Intense Quote')
doc.add_page_break()

# ── 4. POLICY ─────────────────────────────────────────────────────────────────
doc.add_heading('4. Training Management Policy', level=1)

doc.add_heading('4.1 Annual Training Mandate', level=2)
doc.add_paragraph(
    'Every employee is owed a minimum number of training man-hours per financial year by collar. '
    'These targets are policy-configurable per plant; the corporate defaults are:')
table(['Category', 'Default Annual Target', 'Configurable'],
      [['Blue Collared (BC)', '12 man-hours / employee / year', 'Yes — per plant'],
       ['White Collared (WC)', '24 man-hours / employee / year', 'Yes — per plant']],
      widths=[2.2, 2.6, 1.6])
doc.add_paragraph(
    'The financial year runs April to March. Compliance is measured against the active headcount; '
    'see §6 for the exact calculation basis.')

doc.add_heading('4.2 Training Need & Coverage Policy', level=2)
bullets([
    'Training demand is identified through TNI and is the basis for coverage measurement.',
    'Coverage % is measured only over employees nominated in TNI for that programme. Attendance by '
    'a non-nominated employee (a “New Requirement”) is fully credited in seats and man-hours but '
    'does not change coverage %.',
    'Each plant must close its TNI for the year before planning so that coverage targets are real.'])

doc.add_heading('4.3 Data Integrity Controls', level=2)
bullets([
    'Programme Master is the gatekeeper. Programme names in Calendar, 2A and 2C are selected from '
    'the Programme Master only — free-text entry is blocked. This prevents spelling variants from '
    'breaking matching and reporting.',
    'Programme source is exactly one of two values: “TNI Driven” or “New Requirement”.',
    'Audience is locked from TNI. If both BC and WC are nominated → Common; BC only → Blue '
    'Collared; WC only → White Collared. This overrides any manual entry on TNI-driven programmes.',
    'Session Codes are auto-generated and unique; they must never be hand-edited.'])

doc.add_heading('4.4 Conduct, Verification & Approval', level=2)
bullets([
    'A session is recorded as conducted (2C) only after its planned end date and after at least one '
    'attendance record exists.',
    'When a SPOC saves 2C, the session moves to “Awaiting Verification”. It is NOT counted as a '
    'conducted programme until Central L&D verifies it.',
    'Central L&D (or Admin) verification promotes the session to “Conducted”. Only conducted '
    'sessions count in the Monthly Summary and the compliance dashboard.',
    'Rejection returns the session for correction; cancelled / re-scheduled sessions are excluded.'])

doc.add_heading('4.5 Training Effectiveness Policy', level=2)
bullets([
    'Specialized programmes require an effectiveness review.',
    'A review is due 90 days after the conducted date and becomes overdue 30 days after the due '
    'date (i.e. 120 days after conduct).',
    'Effectiveness captures learning gain and on-the-job application (Kirkpatrick L2–L3).'])

doc.add_heading('4.6 Records, Security & Audit', level=2)
bullets([
    'All data is scoped by plant; a SPOC can never see or alter another plant’s data.',
    'Central and Admin accounts require two-factor authentication (2FA).',
    'Idle sessions time out after 30 minutes; first login forces a password change.',
    'Every create / edit / delete and every verification is written to a tamper-evident audit log, '
    'traceable to the user, timestamp (IST), and before/after values.',
    'Timestamps are recorded in IST regardless of server location.'])
doc.add_page_break()

# ── 5. SOP ────────────────────────────────────────────────────────────────────
doc.add_heading('5. Standard Operating Procedure — The Training Cycle', level=1)
doc.add_paragraph(
    'Follow the cycle in order; each step feeds the next. Sequence: '
    'Employees → TNI → Programme Master → Calendar → Conduct (QR Attendance + Feedback) → '
    '2C → 2A → Verification → Effectiveness → Monthly Summary / Dashboard → Export.', style='Intense Quote')

steps = [
    ('5.1 Maintain the Employee Master',
     ['Keep the plant headcount current: add joiners, mark exits, set collar (BC/WC), department, '
      'grade and category.',
      'Only active employees can be marked present or nominated. Exits remove an employee from '
      'future eligibility while preserving their historical records.']),
    ('5.2 Upload & Confirm TNI',
     ['Upload the Training Need Identification sheet. The system fuzzy-matches programme names '
      'against the Programme Master and flags mismatches.',
      'Review the analysis screen and confirm. Correct any unmatched names before saving.']),
    ('5.3 Sync the Programme Master',
     ['Sync the Programme Master from TNI so every demanded programme exists in the canonical list. '
      'Caution: a full sync rebuilds the list from TNI and will remove manually-added New-Requirement '
      'programmes that have no TNI rows — add those back only if intended.']),
    ('5.4 Plan the Calendar',
     ['Create planned sessions. Select the programme from the Programme Master (gatekeeper).',
      'The Session Code is generated automatically; audience is set automatically from TNI.',
      'Set planned dates, planned participants and venue.']),
    ('5.5 Generate the QR (and optional PIN)',
     ['Generate an attendance QR and/or a feedback QR for the planned session. QR generation is '
      'allowed only while the session is “To Be Planned”.',
      'Optionally set a 4-digit session PIN and announce it to participants to restrict scanning to '
      'the room. Print the QR poster.']),
    ('5.6 Conduct — QR Attendance & Feedback (real time)',
     ['During the session, participants scan the attendance QR on any phone (no app, no login), '
      'search their name/code, and mark attendance. The live monitor updates in real time.',
      'After the session, participants scan the feedback QR and rate the programme (9 questions, '
      'available in English and Hindi). Anonymous feedback is de-duplicated per device.',
      'Attendance and feedback must be captured BEFORE recording 2C — recording 2C moves the '
      'session out of the scan-open state.']),
    ('5.7 Record 2C — Programme Details (conducted actuals)',
     ['After the planned end date and with at least one attendee present, record the conducted '
      'actuals: actual dates, faculty, internal/external, actual hours, cost, venue.',
      'Any QR feedback already collected is folded into this record automatically; manually-entered '
      'feedback is never overwritten.',
      'On save the session moves to “Awaiting Verification”. The system flags anomalies '
      '(hours mismatch, low attendance, collar mismatch) for the verifier.']),
    ('5.8 Record 2A — Training Records (attendance back-fill)',
     ['Add any attendees not captured by QR (e.g. manual back-fill or bulk upload). The employee '
      'must be active; the session code auto-fills programme details.',
      'The system prevents double-counting the same person for the same session.']),
    ('5.9 Verify (Central L&D)',
     ['Central L&D reviews the Awaiting-Verification queue, checks the flagged anomalies and '
      'evidence, and approves (→ Conducted) or rejects (→ back for correction).',
      'Only after approval does the session count toward programmes, coverage and compliance.']),
    ('5.10 File the Effectiveness Review',
     ['For Specialized programmes, file the effectiveness review before it becomes overdue '
      '(90 days after conduct; overdue at 120 days). Capture learning gain and application.']),
    ('5.11 Review Monthly Summary & Dashboard',
     ['Review the Monthly Summary (programmes, person-seats, man-hours, coverage %) and the '
      'compliance dashboard (BC% / WC% against the mandate). Numbers update in real time as data '
      'is entered; analytics charts refresh within ~1 minute.',
      'Investigate any plant/department/collar falling behind the mandate.']),
    ('5.12 Export Records',
     ['Export the monthly summary, coverage, attendance and compliance records to Excel for '
      'management review and statutory/audit evidence.']),
]
for title, body in steps:
    doc.add_heading(title, level=2)
    numbered(body)
doc.add_page_break()

# ── 6. CALCULATION BASIS ──────────────────────────────────────────────────────
doc.add_heading('6. Compliance Calculation Basis', level=1)
doc.add_paragraph(
    'The authoritative, line-by-line formulas for all 92 system calculations are maintained in '
    '“TMS — Calculation Logic Reference” (docs/markdown/TMS_Calculation_Logic.md). The policy-level '
    'summary is below.')
table(['Metric', 'Definition', 'Key Rules'],
      [['No. of Programmes', 'Distinct conducted programmes per type & audience.',
        'Counts CONDUCTED sessions only; Awaiting-Verification not counted.'],
       ['Person Seats', 'Count of attendance rows (one per person per session).',
        'BC/WC split by employee collar; financial-year scoped.'],
       ['Man-Hours', 'Sum of training hours across attendance rows.',
        'Financial-year scoped; includes centrally-hosted attendance.'],
       ['Coverage %', 'TNI-nominated employees trained ÷ TNI-nominated, per type & collar.',
        'Denominator = TNI Driven nominations for the FY; non-TNI attendees excluded.'],
       ['Compliance %', 'Man-hours delivered ÷ (active headcount × annual target).',
        'Targets per-plant configurable (BC 12 / WC 24 default).']],
      widths=[1.4, 2.6, 2.4])
doc.add_heading('6.1 Exit Policy (man-hour compliance)', level=2)
doc.add_paragraph(
    'By policy, the man-hours delivered to an employee who later exits remain credited in the '
    'numerator (the training did happen), while the headcount denominator counts only currently '
    'active employees. As a result the compliance gauge can legitimately exceed 100%. This is '
    'intended and must not be read as a data error.')
doc.add_heading('6.2 Cross-Screen Consistency', level=2)
bullets([
    'The compliance gauge, the Monthly Summary compliance card and the Central per-plant view use '
    'the same calculation engine — the same number agrees across screens.',
    'The man-hour drill-in and reports use the same per-plant target as the gauge.',
    'Two coverage figures exist by design: the planning-stage coverage on the Calendar (against raw '
    'demand) and the canonical Coverage % on the Monthly Summary (TNI-Driven, FY-scoped).'])
doc.add_page_break()

# ── 7. CONTROLS & EXCEPTIONS ──────────────────────────────────────────────────
doc.add_heading('7. Controls, Anomalies & Exceptions', level=1)
table(['Control', 'Trigger', 'Action'],
      [['Hours mismatch', '2C hours differ from average 2A hours by more than 25%.',
        'Flagged for verifier; saved with a warning.'],
       ['Low attendance', 'Attendance below 50% of planned participants.',
        'Flagged for verifier.'],
       ['Collar mismatch', 'Attendee collar differs from the session’s locked audience.',
        'Allowed but flagged for Central review.'],
       ['Time vs duration', 'Start/end time inconsistent with stated duration (±tolerance).',
        'Save blocked until corrected.'],
       ['Scan window', 'Attendance/feedback scans accepted only while session is open.',
        'Capture attendance before recording 2C.'],
       ['Programme not in master', 'Programme name not found in Programme Master.',
        'Selection blocked — add to master first.']],
      widths=[1.5, 2.8, 2.1])
doc.add_paragraph(
    'Exceptions to this SOP require written approval from Corporate L&D and must be recorded in the '
    'audit log.')

doc.add_heading('8. Compliance & Enforcement', level=1)
bullets([
    'Each SPOC is accountable for the timeliness and accuracy of their plant’s data.',
    'Central L&D monitors compliance monthly and escalates plants below target.',
    'Deliberate misreporting is a disciplinary matter.',
    'This document is reviewed annually or upon any material change to the training process.'])

footer_pagenum()
doc.save(OUT)
print('Wrote', OUT)
